import io
import locale

from docx import Document
from docx.shared import Cm

from datetime import datetime

locale.setlocale(locale.LC_TIME, 'es_CO')


def createReport(report_data, analysis):

    doc = Document()

    for style in doc.styles:
        try:
            doc.styles[style.name].font.name = 'Arial'
        except:
            None

    # Sección de Encabezado

    header = doc.sections[0].header
    h_content = header.paragraphs[0]

    h_run = h_content.add_run()
    h_run.add_picture('logo.png', width=Cm(1.5))

    text_run = h_content.add_run()
    text_run.text = '\t' + 'Roboat Stingray - Generador de Reportes'

    # Sección de Pié de Página

    bottom = doc.sections[0]
    footer = bottom.footer

    f_content = footer.paragraphs[0]
    f_content.text = '''
    Roboat Stingray Web App
    Universidad de Boyacá © - Todos los derechos reservados (2021) '''

    # Contenido del Documento

    p = doc.add_paragraph('')
    p.add_run(str(datetime.now().strftime("%A, %d de %B de %Y"))).bold = True

    doc.add_heading(report_data['title'], 1) if report_data['title'] else doc.add_heading(
        'Reporte Preliminar de Calidad del Agua', 0)

    doc.add_heading(
        'Análisis preliminar - ' + report_data['sourceType'] + ' ' + report_data['sourceName'] + ', ' + report_data[
            'zoneName'], 2)

    doc.add_paragraph('')

    doc.add_paragraph(
        'Este reporte sobre el análisis preliminar de la calidad del agua es generado en la aplicación web del Sistema Roboat Stingray. Aquí se describen las características de las variables medidas y los datos estadísticos recolectados.')

    doc.add_paragraph('')

    p = doc.add_paragraph('')
    p.add_run('Autor(es): ').bold = True
    p.add_run(report_data['author']) if report_data['author'] else p.add_run('Aquí va el nombre del autor(es)')

    p.add_run('\n\n')

    if report_data['sendTo'] != '':
        p.add_run('Dirigido a: ').bold = True
        p.add_run(report_data['sendTo']) if report_data['sendTo'] else p.add_run(
            'Aquí va la entidad o persona a quien se dirige el reporte')

        p.add_run('\n\n')

    p.add_run('Razón del reporte:\n').bold = True
    p.add_run(report_data['reason']) if report_data['reason'] else p.add_run(
        'Aquí va la descripción del por qué se genera el reporte')

    p.add_run('\n\n')

    if report_data['descrip'] != '':
        p.add_run('Descripción del lugar:\n').bold = True
        p.add_run(report_data['descrip']) if report_data['descrip'] else p.add_run(
            'Aquí va la descripción del lugar de estudio.')

        p.add_run('\n')

    doc.add_heading('Estadísticas de las mediciones:', 2)

    # Tabla
    t = doc.add_table(analysis['C'].shape[0] + 1, analysis['C'].shape[1])
    t.style = 'Colorful Grid Accent 1'

    # add the header rows.
    for j in range(analysis['C'].shape[-1]):
        t.cell(0, j).text = analysis['C'].columns[j]

    # add the rest of the data frame
    for i in range(analysis['C'].shape[0]):
        for j in range(analysis['C'].shape[-1]):
            t.cell(i + 1, j).text = str(analysis['C'].values[i, j])

    doc.add_paragraph('')

    doc.add_heading('Resumen del análisis:', 2)

    doc.add_paragraph(analysis['A'] + ', de las cuales ' + analysis['B'])
    doc.add_paragraph(
        'Hay que recordar que las sondas utilizadas son de marca AtlasScientific, cuyos rangos de medición son:')

    doc.add_paragraph('pH (Lab): desde 0.005 hasta 14.', style='List Bullet')
    doc.add_paragraph('Conductividad (Lab - K 1.0): desde 5 µS/cm hasta 500000 µS/cm.', style='List Bullet')
    doc.add_paragraph('Oxígeno disuelto (Lab): desde 0.1 mg/L hasta 100 mg/L.', style='List Bullet')
    doc.add_paragraph('Temperatura (PT-1000): desde -200ºC hasta 850ºC.', style='List Bullet')

    doc.add_paragraph('Se identificó que el ' + str(
        analysis['F']['% Cumplimiento Total'][1]) + '% del agua de la fuente hídrica, monitoreada en la ubicación ' +
                      report_data[
                          'zoneName'] + ', cumple con las resoluciones ambientales: Resolución 2115 de 2007 y Objetivos de Calidad de Agua - Resolución 3560 de 2015; frente al ' + str(
        analysis['F']['% Cumplimiento Total'][0]) + '% que no lo hace.')

    doc.add_paragraph('Es posible decir entonces que el ' + str(
        analysis['F']['% Cumplimiento Total'][1]) + '% del agua, tiene una alta probabilidad de ser potable.')

    doc.add_heading('Tabla de cumplimiento total:', 2)
    doc.add_paragraph(
        'Esta tabla muestra el porcentaje de cumplimiento de ambas resoluciones sobre el total de tramos analizados, a partir del recorrido realizado por el sistema Roboat Stingray en la fuente hídrica.')

    t = doc.add_table(analysis['F'].shape[0] + 1, analysis['F'].shape[1])
    t.style = 'Colorful Grid Accent 1'

    # add the header rows.
    for j in range(analysis['F'].shape[-1]):
        t.cell(0, j).text = analysis['F'].columns[j]

    # add the rest of the data frame
    for i in range(analysis['F'].shape[0]):
        for j in range(analysis['F'].shape[-1]):
            t.cell(i + 1, j).text = str(analysis['F'].values[i, j])

    doc.add_paragraph('')

    doc.add_heading('Tabla de cumplimiento por variable:', 2)
    doc.add_paragraph(
        'Esta tabla muestra el porcentaje de cumplimiento las resoluciones para cada variable, sobre el total de mediciones analizadas. Hay que tener en cuenta que la Resolución 2115 evalúa el pH y la Conductividad, y que los Objetivos de Calidad de Agua evalúan el oxígeno disuelto; para la temperatura, el rango tolerable se programó entre los 5ºC hasta los 30ºC.')

    t = doc.add_table(analysis['E'].shape[0] + 1, analysis['E'].shape[1])
    t.style = 'Colorful Grid Accent 1'

    # add the header rows.
    for j in range(analysis['E'].shape[-1]):
        t.cell(0, j).text = analysis['E'].columns[j]

    # add the rest of the data frame
    for i in range(analysis['E'].shape[0]):
        for j in range(analysis['E'].shape[-1]):
            t.cell(i + 1, j).text = str(analysis['E'].values[i, j])

    doc.add_paragraph('')

    doc.add_heading('Mapa del recorrido:', 2)
    doc.add_paragraph(
        'Inserte aquí su captura de pantalla del recorrido realizado por el catamarán del sistema para añadir más información valiosa a si reporte.')

    # Guardado del archivo en memoria

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return file_stream
